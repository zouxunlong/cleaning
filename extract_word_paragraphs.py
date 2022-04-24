import time
import pycld2 as cld2
import cld3
import fasttext
from docx import Document
import os
import re
import string
from docx.oxml.shared import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from googletrans import Translator

translator = Translator()
start_time = time.time()


def GetParagraphRuns(paragraph):
    def _get(node, parent):
        for child in node:
            if child.tag == qn('w:r'):
                yield Run(child, parent)
            if child.tag == qn('w:hyperlink'):
                yield from _get(child, parent)
    return list(_get(paragraph._element, paragraph))


Paragraph.runs = property(lambda self: GetParagraphRuns(self))

model_fasttext = fasttext.load_model('../model/lid.176.bin')


def allocate_text_by_lang(texts):
    texts_en = []
    texts_ms = []
    texts_zh = []
    texts_ta = []

    lang_by_cld2 = ""
    lang_by_cld3 = ""
    lang_by_fasttext = ""
    lang_by_google = ""

    for text in texts:

        text = re.sub("^[a-zA-Z0-9]{0,2}\.\s?", "", text).strip()

        trimed_text = re.sub(
            "\w+@\S+\s?|http\S*\s?|www\.\S*\s?|[a-z\.]*\.sg\S*\s?", "", text)

        text_for_lang_detect = trimed_text.translate(
            str.maketrans('', '', string.punctuation)).strip().lower()

        if text_for_lang_detect:

            lang_by_cld2 = cld2.detect(text_for_lang_detect)[2][0][1]
            lang_by_cld3 = cld3.get_language(text_for_lang_detect)[0]
            lang_by_fasttext = model_fasttext.predict(
                text_for_lang_detect)[0][0]

            if len(text_for_lang_detect) <= 6:
                lang_by_google = translator.detect(text_for_lang_detect).lang
            else:
                lang_by_google = ""

        if lang_by_google[:2] == "en" or lang_by_cld2 == "en" or lang_by_fasttext == "__label__en":
            texts_en.append(text)
        elif lang_by_google[:2] == "ms" or lang_by_cld2 in ["ms", "id"] or lang_by_fasttext in ["__label__ms", "__label__id"] or lang_by_cld3 in ["ms"]:
            texts_ms.append(text)
        elif lang_by_google[:2] == "zh" or lang_by_cld2 in ["zh", "ja"] or lang_by_fasttext in ["__label__zh", "__label__ja"]:
            text = text.replace(" ", "")
            texts_zh.append(text)
        elif lang_by_google[:2] == "ms" or lang_by_cld2 == "ta" or lang_by_fasttext == "__label__ta":
            texts_ta.append(text)

    return texts_en, texts_ms, texts_zh, texts_ta


def texts_from_tables(tables):
    def yield_texts(_tables):
        for table in _tables:
            for row in table.rows:
                for cell in row.cells:
                    texts = [text.strip() for text in cell.text.split('\n')]
                    for text in texts:
                        yield text
                    if cell.tables:
                        yield from yield_texts(cell.tables)
    return [text for text in list(yield_texts(tables)) if text]


def texts_from_paragraphs(paragraphs):
    texts = [text.strip() for p in paragraphs for text in p.text.split('\n')]
    return [text for text in texts if text]


def texts_from_textboxs(root_element):
    textbox_elements = root_element.xpath('.//w:drawing//w:txbxContent')
    texts = [" ".join(" ".join(textbox_element.xpath(".//text()")).split())
             for textbox_element in textbox_elements]
    return [text for text in texts if text]


def extract(filepath):
    wordDoc = Document(filepath)

    texts = []

    texts.extend(texts_from_tables(wordDoc.tables))
    texts.extend(texts_from_paragraphs(wordDoc.paragraphs))

    for section in wordDoc.sections:
        header = section.header
        footer = section.footer
        texts.extend(texts_from_paragraphs(header.paragraphs))
        texts.extend(texts_from_paragraphs(footer.paragraphs))

    texts.extend(texts_from_textboxs(wordDoc.element))

    texts_en, texts_ms, texts_zh, texts_ta = allocate_text_by_lang(texts)

    print("texts_en number:{}".format(len(texts_en)))
    print("texts_ms number:{}".format(len(texts_ms)))
    print("texts_zh number:{}".format(len(texts_zh)))
    print("texts_ta number:{}".format(len(texts_ta)))

    with open('./sentences.en', 'w', encoding='utf8') as fOut:
        for sentence in texts_en:
            fOut.write("{}\n".format(sentence.replace("|", " ")))

    with open('./sentences.zh', 'w', encoding='utf8') as fOut:
        for sentence in texts_zh:
            fOut.write("{}\n".format(sentence.replace("|", " ")))

    with open('./sentences.ms', 'w', encoding='utf8') as fOut:
        for sentence in texts_ms:
            fOut.write("{}\n".format(sentence.replace("|", " ")))

    with open('./sentences.ta', 'w', encoding='utf8') as fOut:
        for sentence in texts_ta:
            fOut.write("{}\n".format(sentence.replace("|", " ")))


extract("LED Screen_Video Loop_User Guide_Master_0709  (ECMT).docx")

print("--- %s seconds ---" % (time.time() - start_time))

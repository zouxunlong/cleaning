from docx import Document
import os
import re
import string
from docx.oxml.shared import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from googletrans import Translator

translator = Translator()


def get_paragraph_runs(paragraph):
    def _get(node, parent):
        for child in node:
            if child.tag == qn('w:r'):
                yield Run(child, parent)
            if child.tag == qn('w:hyperlink'):
                yield from _get(child, parent)
    return list(_get(paragraph._element, paragraph))


Paragraph.runs = property(lambda self: get_paragraph_runs(self))


def texts_from_tables(tables):
    def yield_texts(_tables):
        for table in _tables:
            for column in table.columns:
                for cell in column.cells:
                    texts = [text.strip() for text in cell.text.split('\n')]
                    for text in texts:
                        yield text
                    if cell.tables:
                        yield from yield_texts(cell.tables)
    return [text for text in list(yield_texts(tables)) if text]


def texts_from_paragraphs(paragraphs):
    texts = [text.strip() for p in paragraphs for text in p.text.split('\n')]
    return [text for text in texts if text]


def texts_from_textboxs(root_element):
    textbox_elements = root_element.xpath('.//w:drawing//w:txbxContent')
    texts = [" ".join(" ".join(textbox_element.xpath(".//text()")).split())
             for textbox_element in textbox_elements]
    return [text for text in texts if text]


def extend_texts_from_docx(docx_path, texts):

    wordDoc = Document(docx_path)
    for section in wordDoc.sections:
        header = section.header
        footer = section.footer
        texts.extend(texts_from_paragraphs(header.paragraphs))
        texts.extend(texts_from_paragraphs(footer.paragraphs))
    texts.extend(texts_from_tables(wordDoc.tables))
    texts.extend(texts_from_paragraphs(wordDoc.paragraphs))
    texts.extend(texts_from_textboxs(wordDoc.element))


def extend_texts_from_file(file_path, texts):

    if file.endswith('.docx'):
        extend_texts_from_docx(file_path, texts)


rootdir = './'


for root, dirs, files in os.walk(rootdir):
    files.sort()
    texts = []
    for i, file in enumerate(files):

        file_path = os.path.join(root, file)

        extend_texts_from_file(file_path, texts)

        if texts:

            print(texts)

            texts.clear()


print('finished translated sentences mining')


